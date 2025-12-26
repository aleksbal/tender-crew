from __future__ import annotations

import os
import re
from dataclasses import dataclass
from typing import List, Tuple, Optional, Dict

import fitz  # pymupdf
from docx import Document


@dataclass
class ExtractDiagnostics:
    file_type: str
    pages: int = 0
    rejected_scanned_pdf: bool = False
    multi_column_pages: int = 0
    empty_text_pages: int = 0


SECTIONY_RE = re.compile(
    r"^\s*(profil|zusammenfassung|skills|kenntnisse|technologien|"
    r"berufserfahrung|erfahrung|projekte|ausbildung|zertifikate|sprachen|"
    r"experience|projects|education|certificates|languages)\s*$",
    re.IGNORECASE,
)


def extract_text(path: str) -> Tuple[str, ExtractDiagnostics]:
    ext = os.path.splitext(path.lower())[1]
    if ext == ".docx":
        text = _docx_to_text(path)
        return text, ExtractDiagnostics(file_type="docx", pages=1)
    if ext == ".pdf":
        text, diag = _pdf_to_text_markdownish(path)
        return text, diag
    raise ValueError(f"Unsupported file type: {ext}")


# ---------------- DOCX ----------------

def _docx_to_text(path: str) -> str:
    doc = Document(path)
    lines: List[str] = []

    # Paragraphs (keeps bullets as plain text; good enough for LLM extraction)
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    # Tables → flatten rows
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines).strip() + "\n"


# ---------------- PDF ----------------

def _pdf_to_text_markdownish(path: str) -> Tuple[str, ExtractDiagnostics]:
    doc = fitz.open(path)
    diag = ExtractDiagnostics(file_type="pdf", pages=doc.page_count)
    out_lines: List[str] = []

    for pno in range(doc.page_count):
        page = doc.load_page(pno)

        blocks = page.get_text("blocks")  # (x0, y0, x1, y1, "text", block_no, block_type)
        # Keep only text blocks with some content
        text_blocks = [
            (x0, y0, x1, y1, txt)
            for (x0, y0, x1, y1, txt, *_rest) in blocks
            if isinstance(txt, str) and txt.strip()
        ]

        if not text_blocks:
            # No selectable text -> likely scanned or weird embedded images → reject by your policy
            diag.rejected_scanned_pdf = True
            diag.empty_text_pages += 1
            raise ValueError(
                f"PDF page {pno+1} has no extractable text (scanned/bitmap?). Rejecting by policy."
            )

        # Decide whether page is likely multi-column
        is_multi = _is_likely_two_column(text_blocks, page_width=page.rect.width)
        if is_multi:
            diag.multi_column_pages += 1

        out_lines.append(f"# Page {pno+1}")
        out_lines.append("")  # blank

        # Order blocks into reading order (single or two-column)
        ordered_texts = _order_blocks_reading_order(text_blocks, two_column=is_multi, page_width=page.rect.width)

        # Light cleanup + add structure cues
        out_lines.extend(_normalize_block_texts(ordered_texts))
        out_lines.append("")  # page break spacer
        out_lines.append("---")
        out_lines.append("")

    return "\n".join(out_lines).strip() + "\n", diag


def _is_likely_two_column(text_blocks: List[Tuple[float, float, float, float, str]], page_width: float) -> bool:
    """
    Simple heuristic:
    - compute block centers; if many blocks cluster clearly into left/right halves,
      treat as two-column.
    """
    centers = [(x0 + x1) / 2 for (x0, _y0, x1, _y1, _t) in text_blocks]
    mid = page_width / 2

    left = sum(1 for c in centers if c < mid * 0.95)
    right = sum(1 for c in centers if c > mid * 1.05)

    # require both sides to have enough blocks
    total = len(centers)
    return left >= max(4, total * 0.25) and right >= max(4, total * 0.25)


def _order_blocks_reading_order(
    text_blocks: List[Tuple[float, float, float, float, str]],
    two_column: bool,
    page_width: float,
) -> List[str]:
    """
    For single-column:
      sort by y0 then x0
    For two-column:
      split into left/right by x center, sort each by y0, then left column first, then right.
    """
    if not two_column:
        blocks_sorted = sorted(text_blocks, key=lambda b: (b[1], b[0]))  # y0, x0
        return [b[4] for b in blocks_sorted]

    mid = page_width / 2
    left_blocks = []
    right_blocks = []
    for (x0, y0, x1, y1, t) in text_blocks:
        cx = (x0 + x1) / 2
        if cx <= mid:
            left_blocks.append((x0, y0, x1, y1, t))
        else:
            right_blocks.append((x0, y0, x1, y1, t))

    left_sorted = sorted(left_blocks, key=lambda b: (b[1], b[0]))
    right_sorted = sorted(right_blocks, key=lambda b: (b[1], b[0]))

    return [b[4] for b in left_sorted] + [b[4] for b in right_sorted]


def _normalize_block_texts(block_texts: List[str]) -> List[str]:
    """
    Turn blocks into lines, preserve bullets, add blank lines between logical segments.
    Adds lightweight heading markers if a line looks like a section title.
    """
    out: List[str] = []
    for block in block_texts:
        # split to lines, strip excessive whitespace
        lines = [re.sub(r"\s+", " ", ln).strip() for ln in block.splitlines()]
        lines = [ln for ln in lines if ln]

        for ln in lines:
            # Promote obvious section headers
            if SECTIONY_RE.match(ln):
                out.append(f"## {ln}")
            else:
                out.append(ln)

        # spacer between blocks
        out.append("")
    return out


def _normalize_block(block: str) -> List[str]:
    """
    Normalize a single text block into a list of lines (preserves bullets,
    promotes section-like lines to '## ' headings) — used for structured output.
    """
    lines = [re.sub(r"\s+", " ", ln).strip() for ln in block.splitlines()]
    lines = [ln for ln in lines if ln]
    out: List[str] = []
    for ln in lines:
        if SECTIONY_RE.match(ln):
            out.append(f"## {ln}")
        else:
            out.append(ln)
    return out


def _order_blocks_reading_order_blocks(
    text_blocks: List[Tuple[float, float, float, float, str]],
    two_column: bool,
    page_width: float,
) -> List[Tuple[float, float, float, float, str]]:
    """
    Like `_order_blocks_reading_order` but returns the full block tuples
    including their bboxes so callers can emit structured output with bboxes.
    """
    if not two_column:
        blocks_sorted = sorted(text_blocks, key=lambda b: (b[1], b[0]))  # y0, x0
        return blocks_sorted

    mid = page_width / 2
    left_blocks = []
    right_blocks = []
    for (x0, y0, x1, y1, t) in text_blocks:
        cx = (x0 + x1) / 2
        if cx <= mid:
            left_blocks.append((x0, y0, x1, y1, t))
        else:
            right_blocks.append((x0, y0, x1, y1, t))

    left_sorted = sorted(left_blocks, key=lambda b: (b[1], b[0]))
    right_sorted = sorted(right_blocks, key=lambda b: (b[1], b[0]))

    return left_sorted + right_sorted


def _pdf_to_structured(path: str) -> Tuple[Dict, ExtractDiagnostics]:
    """
    Produce structured JSON-like dict for a PDF:
      { file_type: 'pdf', pages: [ { page_number, width, height, page_text, blocks: [ {bbox, text, lines:[{text,char_start,char_end}] } ] } ] }

    Returns (structured_dict, diagnostics).
    """
    doc = fitz.open(path)
    diag = ExtractDiagnostics(file_type="pdf", pages=doc.page_count)
    pages_out: List[Dict] = []

    for pno in range(doc.page_count):
        page = doc.load_page(pno)

        blocks = page.get_text("blocks")
        text_blocks = [
            (x0, y0, x1, y1, txt)
            for (x0, y0, x1, y1, txt, *_rest) in blocks
            if isinstance(txt, str) and txt.strip()
        ]

        if not text_blocks:
            diag.rejected_scanned_pdf = True
            diag.empty_text_pages += 1
            raise ValueError(
                f"PDF page {pno+1} has no extractable text (scanned/bitmap?). Rejecting by policy."
            )

        is_multi = _is_likely_two_column(text_blocks, page_width=page.rect.width)
        if is_multi:
            diag.multi_column_pages += 1

        ordered_blocks = _order_blocks_reading_order_blocks(text_blocks, two_column=is_multi, page_width=page.rect.width)

        page_text = ""
        blocks_out: List[Dict] = []

        for (x0, y0, x1, y1, txt) in ordered_blocks:
            norm_lines = _normalize_block(txt)
            block_text = "\n".join(norm_lines)

            if page_text:
                sep = "\n\n"
                abs_block_start = len(page_text) + len(sep)
                page_text = page_text + sep + block_text
            else:
                abs_block_start = 0
                page_text = block_text

            # compute per-line spans relative to page_text
            line_spans: List[Dict] = []
            pos = 0
            for i, ln in enumerate(norm_lines):
                ln_start = pos
                ln_end = ln_start + len(ln)
                line_spans.append({
                    "text": ln,
                    "char_start": abs_block_start + ln_start,
                    "char_end": abs_block_start + ln_end,
                })
                pos = ln_end + 1  # account for the '\n' between lines in block_text

            blocks_out.append({
                "bbox": [x0, y0, x1, y1],
                "text": block_text,
                "lines": line_spans,
            })

        pages_out.append({
            "page_number": pno + 1,
            "width": page.rect.width,
            "height": page.rect.height,
            "page_text": page_text,
            "blocks": blocks_out,
        })

    return {"file_type": "pdf", "pages": pages_out}, diag


def extract_text_structured(path: str) -> Tuple[Dict, ExtractDiagnostics]:
    """
    New API: return a structured representation (dict) and diagnostics for a file.
    For PDFs returns the full pages→blocks→lines structure; for DOCX returns a single-page representation
    with paragraph/table blocks.
    """
    ext = os.path.splitext(path.lower())[1]
    if ext == ".pdf":
        return _pdf_to_structured(path)
    if ext == ".docx":
        # Build a simple single-page structured view for DOCX
        doc = Document(path)
        diag = ExtractDiagnostics(file_type="docx", pages=1)
        blocks_out: List[Dict] = []

        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                norm_lines = _normalize_block(t)
                blocks_out.append({
                    "bbox": None,
                    "text": "\n".join(norm_lines),
                    "lines": [{"text": ln, "char_start": None, "char_end": None} for ln in norm_lines],
                })

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    t = " | ".join(cells)
                    norm_lines = _normalize_block(t)
                    blocks_out.append({
                        "bbox": None,
                        "text": "\n".join(norm_lines),
                        "lines": [{"text": ln, "char_start": None, "char_end": None} for ln in norm_lines],
                    })

        page_text = "\n\n".join(b["text"] for b in blocks_out)
        page = {"page_number": 1, "width": None, "height": None, "page_text": page_text, "blocks": blocks_out}
        return {"file_type": "docx", "pages": [page]}, diag

    raise ValueError(f"Unsupported file type: {ext}")
